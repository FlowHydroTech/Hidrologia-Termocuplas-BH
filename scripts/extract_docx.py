from pathlib import Path
from docx import Document

in_path = Path('doc/Manual_Completo_VFLUX2_SebastianErazo.docx')
out_dir = Path('doc/extracted')
out_dir.mkdir(parents=True, exist_ok=True)
out_path = out_dir / 'Manual_Completo_VFLUX2_SebastianErazo.txt'

doc = Document(in_path)
lines = []
for p in doc.paragraphs:
    text = p.text.strip()
    if text:
        lines.append(text)

# Also extract tables if any
for table in doc.tables:
    for row in table.rows:
        row_text = '\t'.join(cell.text.strip() for cell in row.cells)
        if row_text.strip():
            lines.append(row_text)

out_path.write_text('\n'.join(lines), encoding='utf-8')
print('Texto extraído en:', out_path)
